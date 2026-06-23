import logging
import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class Document:
    id: str
    filename: str
    content: str
    file_type: str
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        self.metadata.setdefault("char_count", len(self.content))
        self.metadata.setdefault("word_count", len(self.content.split()))


def _generate_doc_id(filename: str, content: str) -> str:
    payload = f"{filename}{content[:200]}"
    return hashlib.md5(payload.encode()).hexdigest()[:12]


def _load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber: pip install pdfplumber")

    pages = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages.append(f"[Page {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def _load_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


_LOADERS = {
    ".txt": _load_txt,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".md": _load_markdown,
}

SUPPORTED_FORMATS = list(_LOADERS.keys())


def load_document(file_path: str | Path) -> Document:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in _LOADERS:
        raise ValueError(f"Unsupported format '{suffix}'. Supported: {SUPPORTED_FORMATS}")

    loader = _LOADERS[suffix]
    content = loader(path)

    if not content.strip():
        raise ValueError(f"No text could be extracted from '{path.name}'")

    doc_id = _generate_doc_id(path.name, content)
    logger.info(f"Loaded '{path.name}' → {len(content):,} chars, id={doc_id}")

    return Document(
        id=doc_id,
        filename=path.name,
        content=content,
        file_type=suffix.lstrip("."),
        metadata={"source_path": str(path), "file_size_bytes": path.stat().st_size},
    )


def load_documents_from_directory(directory: str | Path) -> list[Document]:
    directory = Path(directory)
    if not directory.is_dir():
        raise NotADirectoryError(f"Not a directory: {directory}")

    docs = []
    for path in sorted(directory.iterdir()):
        if path.suffix.lower() in _LOADERS and not path.name.startswith("."):
            try:
                docs.append(load_document(path))
            except Exception as e:
                logger.warning(f"Skipping '{path.name}': {e}")

    logger.info(f"Loaded {len(docs)} documents from '{directory}'")
    return docs
